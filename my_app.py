from docx import Document
from docx.shared import Inches
import pyttsx3

# função de fala 
def speak(text):
    pyttsx3.speak(text)

# Cria o documento
document = Document()

# Detalhes sobre a pessoa
name = input ('Qual é o seu nome? ')
speak('Olá ' + name + ' Como é que estás hoje?')
speak('Qual é o seu numero de telefone?')
phone_number =  input ('Qual é o seu numero de telefone? ')
email = input ('Qual é o seu email? ')

# adiciona uma imagem ao documento e pode por a largura com width
document.add_picture('profilepic.jpg', width = Inches(2.0))

# Escreve texto no documento e adiciona um paragrafo 
document.add_paragraph(name + ' | ' + phone_number + ' | ' + email)

# Sobre mim
# Cria um subtitulo com heading
document.add_heading('Sobre mim')
about_me = input ('Fale me sobre si ')
document.add_paragraph(about_me)

# Experiencias de trabalho
# Cria um subtitulo com heading
document.add_heading('Experiencias de trabalho')
p = document.add_paragraph()

company = input ('Nome da companhia: ')
from_date = input ('Insira o ano que iniciou: ')
to_date = input ('Insira o ano que terminou ou o currente: ')

p.add_run(company + ' ').bold = True
p.add_run(from_date + '-' + to_date + '\n').italic = True

Experiencias_de_trabalho = input('Descreva a sua experiencia na ' + company + ' ')
p.add_run(Experiencias_de_trabalho)

# Mais Experiencias de trabalho
while True:
        mais_experiencias_de_trabalho = input('Tem mais experiencias de trabalho? Sim ou Não ')
        
        if mais_experiencias_de_trabalho.lower() == 'sim':
            p = document.add_paragraph()

            company = input ('Nome da companhia: ')
            from_date = input ('Insira o ano que iniciou: ')
            to_date = input ('Insira o ano que terminou ou o currente: ')

            p.add_run(company + ' ').bold = True
            p.add_run(from_date + '-' + to_date + '\n').italic = True

            Experiencias_de_trabalho = input('Descreva a sua experiencia na ' + company + ' ')
            p.add_run(Experiencias_de_trabalho)
        else:
            break


# Habilidades de Curriculo
document.add_heading('Habilidades de Curriculo')
habilidades = input('Insira a habilidade ')
p = document.add_paragraph(habilidades)
p.style = 'List Bullet'

while True:
        tem_mais_habilidades = input('Tem mais habilidades de curriculo? Sim ou Não ')
        if tem_mais_habilidades.lower() == 'sim':
            habilidades = input('Insira a habilidade ')
            p = document.add_paragraph(habilidades)
            p.style = 'List Bullet'
        else:
            break


# footer ou Rodapé
# Adiciona uma secção chamada rodapé
section = document.sections[0]
footer = section.footer
p = footer.paragraphs[0]
p.text = "CV gerado usando amigoscode"
            


# Grava o documento
document.save('cv.docx')
